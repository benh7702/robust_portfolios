from __future__ import annotations

from pathlib import Path
import json
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
RESULTS = ROOT / "results"
FIG = RESULTS / "figures"
OUT_DOCX = ROOT / "report" / "Ben_Heskin_Robust_Portfolio_Empirical_Report.docx"

summary = pd.read_csv(RESULTS / "performance_summary.csv", index_col=0)
bootstrap = pd.read_csv(RESULTS / "bootstrap_vs_equal_weight.csv", index_col=0)
regimes = pd.read_csv(RESULTS / "regime_summary.csv")
stress = pd.read_csv(RESULTS / "stress_period_summary.csv")
hp = pd.read_csv(RESULTS / "selected_hyperparameters.csv")
metadata = json.loads((RESULTS / "run_metadata.json").read_text())

BLUE = "17365D"
MID = "365F91"
LIGHT = "DCE6F1"
PALE = "F4F7FA"
DARK = "1F2933"
GREY = "52606D"
WHITE = "FFFFFF"


def set_font(run, name="Liberation Serif", size=9.4, bold=False, italic=False, color=DARK):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rpr.get_or_add_rFonts().set(qn("w:ascii"), name)
    rpr.get_or_add_rFonts().set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=55, start=65, bottom=55, end=65):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = paragraph.add_run("Page ")
    set_font(r, "Liberation Sans", 8, color=GREY)
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = " PAGE "
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    rr = OxmlElement("w:r")
    rr.extend([begin, instr, end])
    paragraph._p.append(rr)


def add_rule(paragraph, color="9FB3C8", size="6"):
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    ppr.append(pbdr)


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(0 if level == 1 else 3)
    p.paragraph_format.space_after = Pt(4 if level == 1 else 2)
    r = p.add_run(text)
    if level == 1:
        set_font(r, "Liberation Sans", 15, bold=True, color=BLUE)
        add_rule(p)
    else:
        set_font(r, "Liberation Sans", 11, bold=True, color=MID)
    return p


def body(doc, text, after=3.0, size=9.4, justify=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.04
    r = p.add_run(text)
    set_font(r, size=size)
    return p


def bullet(doc, text, size=9.2, after=1.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.0
    set_font(p.add_run(text), size=size)
    return p


def equation(doc, text, caption=None):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    t.columns[0].width = Inches(6.65)
    c = t.cell(0, 0)
    shade(c, PALE)
    cell_margins(c, 70, 110, 70, 110)
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    set_font(p.add_run(text), size=10.2, italic=True, color=BLUE)
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(2)
        set_font(cp.add_run(caption), "Liberation Sans", 7.8, italic=True, color=GREY)


def callout(doc, title, text, fill="EAF0F6"):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    t.columns[0].width = Inches(6.65)
    c = t.cell(0, 0)
    shade(c, fill)
    cell_margins(c, 80, 110, 80, 110)
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_font(p.add_run(title + "  "), "Liberation Sans", 9.0, bold=True, color=BLUE)
    set_font(p.add_run(text), size=9.0)


def table(doc, headers, rows, widths=None, font_size=8.0):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    t.autofit = False
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        shade(c, BLUE)
        cell_margins(c)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(h), "Liberation Sans", font_size, bold=True, color=WHITE)
        if widths: c.width = Inches(widths[j])
    for i, row in enumerate(rows):
        cells = t.add_row().cells
        for j, val in enumerate(row):
            c = cells[j]
            if i % 2: shade(c, "F7F9FB")
            cell_margins(c)
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 0.95
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
            set_font(p.add_run(str(val)), size=font_size, bold=(j == 0))
            if widths: c.width = Inches(widths[j])
    return t


def figure(doc, path, width, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    p.add_run().add_picture(str(path), width=Inches(width))
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(2)
    set_font(cp.add_run(caption), "Liberation Sans", 7.9, italic=True, color=GREY)


def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def pct(x, d=1):
    return f"{100*x:.{d}f}%"


def num(x, d=2):
    return f"{x:.{d}f}"


doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.58)
sec.bottom_margin = Inches(0.55)
sec.left_margin = Inches(0.68)
sec.right_margin = Inches(0.68)
sec.header_distance = Inches(0.22)
sec.footer_distance = Inches(0.22)

styles = doc.styles
styles["Normal"].font.name = "Liberation Serif"
styles["Normal"].font.size = Pt(9.4)
styles["List Bullet"].font.name = "Liberation Serif"
styles["List Bullet"].font.size = Pt(9.2)

header = sec.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_font(header.add_run("BEN HESKIN  |  ROBUST PORTFOLIO CONSTRUCTION"), "Liberation Sans", 7.6, bold=True, color=GREY)
add_rule(header, color="D5DEE8", size="4")
add_page_number(sec.footer.paragraphs[0])

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(7)
set_font(p.add_run("ESTIMATION ERROR IN\nPORTFOLIO OPTIMISATION"), "Liberation Sans", 23, bold=True, color=BLUE)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(12)
set_font(p.add_run("Bayesian Shrinkage, CVaR and Wasserstein Distributional Robustness"), "Liberation Sans", 12.5, italic=True, color=MID)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(15)
set_font(p.add_run("Ben Heskin  |  Independent Quantitative Finance Research Project  |  August 2026"), "Liberation Sans", 9, color=GREY)
heading(doc, "Abstract", 2)
body(doc, "I compare five monthly portfolio rules across twelve value-weighted US industry portfolios: equal weighting, sample mean-variance optimisation, empirical-Bayes/Ledoit-Wolf mean-variance optimisation, historical 95% Conditional Value-at-Risk (CVaR) optimisation, and Wasserstein robust CVaR. The study uses a nested 60-month walk-forward design from January 2015 to September 2025. Hyperparameters are selected before each test month, and each strategy pays 10 basis points per unit of one-way turnover.", after=4)
heading(doc, "Principal findings", 2)
for text in [
    f"The 1/N benchmark achieved the highest compound return, {pct(summary.loc['Equal Weight','Annualized Return'],2)} per year, and terminal wealth of {num(summary.loc['Equal Weight','Terminal Wealth'])} from an initial dollar.",
    f"Sample MVO was fragile: it returned {pct(summary.loc['Sample MVO','Annualized Return'],2)}, suffered a {pct(summary.loc['Sample MVO','Maximum Drawdown'],1)} maximum drawdown, and turned over {pct(summary.loc['Sample MVO','Average Monthly Turnover'],1)} per month.",
    f"Shrinkage MVO cut annualised volatility to {pct(summary.loc['Shrinkage MVO','Annualized Volatility'],1)} and improved maximum drawdown to {pct(summary.loc['Shrinkage MVO','Maximum Drawdown'],1)}, but did not recover the benchmark's return.",
    f"Wasserstein robust CVaR produced the highest zero-rate Sharpe ratio among the optimisation methods ({num(summary.loc['Wasserstein Robust CVaR','Zero-Rate Sharpe'],3)}) and a more diversified allocation, but its Sharpe difference from 1/N was not statistically distinguishable from zero.",
]: bullet(doc, text)
callout(doc, "Interpretation", "The modern methods mainly improved risk control and stability. They did not establish a statistically reliable return advantage over naive diversification.")
page_break(doc)

heading(doc, "1. Research question and motivation")
body(doc, "Markowitz (1952) frames investment choice as a trade-off between expected return and covariance risk. In practice, both inputs are estimated with error, and optimisation can turn small errors into large changes in portfolio weights.")
body(doc, "The comparison adds empirical-Bayes mean shrinkage, covariance shrinkage, a lower-tail CVaR objective, and Wasserstein robustness to the classical model.")
heading(doc, "1.1 Testable research question", 2)
callout(doc, "Research question", "When the data-generating distribution is estimated from a short rolling history, do shrinkage and distributionally robust tail-risk models generate better out-of-sample portfolios than sample Markowitz optimisation and equal weighting?")
heading(doc, "1.2 Prior evidence", 2)
body(doc, "Ledoit and Wolf (2004) show why a convex combination of the sample covariance matrix and a structured target can be better conditioned than the raw sample estimator. Rockafellar and Uryasev (2000) provide a linear-programming representation of CVaR, making tail-risk optimisation computationally practical. Mohajerin Esfahani and Kuhn (2018) show that data-driven optimisation over Wasserstein ambiguity sets can often be reformulated as tractable convex programmes. DeMiguel, Garlappi and Uppal (2009), however, demonstrate that 1/N is difficult to beat out of sample because estimation error can offset the theoretical gains from optimisation.")
heading(doc, "1.3 Pre-specified hypotheses", 2)
for text in [
    "H1: sample MVO will exhibit the highest concentration and turnover.",
    "H2: mean and covariance shrinkage will reduce realised volatility and drawdown relative to sample MVO.",
    "H3: CVaR-based portfolios will improve lower-tail risk relative to variance-based portfolios.",
    "H4: Wasserstein robustness will reduce concentration relative to historical CVaR.",
    "H5: no advanced model will show a statistically reliable return advantage over 1/N after costs.",
]: bullet(doc, text)
page_break(doc)

heading(doc, "2. Data and experimental protocol")
body(doc, "The asset universe is the twelve value-weighted US industry portfolios from the Kenneth R. French Data Library: non-durables, durables, manufacturing, energy, chemicals, business equipment, telecommunications, utilities, shops, health, money and other. The bundled research vintage contains 189 contiguous monthly observations from January 2010 to September 2025. Returns are total percentage returns converted to decimals.")
rows = [
    ("Universe", "12 US industry portfolios"),
    ("Data frequency", "Monthly"),
    ("Input sample", "2010-01 to 2025-09"),
    ("Out-of-sample period", "2015-01 to 2025-09 (129 months)"),
    ("Estimation window", "Rolling 60 months"),
    ("Internal validation", "Final 12 months of each estimation window"),
    ("Portfolio constraints", "Long-only; sum to 1; 35% cap per industry"),
    ("Trading cost", "10 bps per unit of one-way turnover"),
    ("Tail probability", "95% CVaR"),
]
table(doc, ["Design element", "Specification"], rows, widths=[2.25, 4.4], font_size=8.4)
heading(doc, "2.1 Nested walk-forward procedure", 2)
for text in [
    "At each test month, take only the immediately preceding 60 observations.",
    "Use the first 48 months to fit candidate models and the final 12 months to select MVO risk aversion or the Wasserstein radius by realised validation Sharpe.",
    "Refit the selected specification on all 60 observations, fix the portfolio weights, and then reveal the next monthly return.",
    "Calculate turnover from drifted pre-trade weights, deduct costs, and roll forward by one month.",
]: bullet(doc, text)
heading(doc, "2.2 Bias controls", 2)
body(doc, "The test month is never used for model selection. Parameters are not tuned on the final sample, and all reported returns are net of trading costs. Equal weight uses the same cost convention. The report uses a fixed French data vintage because CRSP revisions can change historical results.")
callout(doc, "Measurement convention", "The reported Sharpe statistic assumes a zero monthly risk-free rate and is labelled 'zero-rate Sharpe'. This avoids mixing the portfolio comparison with a separately sourced interest-rate series.")
page_break(doc)

heading(doc, "3. Statistical estimators and allocation rules")
heading(doc, "3.1 Classical sample mean-variance optimisation", 2)
body(doc, "For monthly return vector r_t, the classical model estimates the sample mean mu_hat and covariance Sigma_hat and solves a long-only quadratic programme. Gamma is selected in the internal validation period.")
equation(doc, "min_w  0.5 gamma w' Sigma_hat w - mu_hat' w", "Subject to 1'w = 1 and 0 <= w_i <= 0.35.")
body(doc, "Expected-return errors enter linearly, while covariance inversion and optimisation can amplify noise. The cap prevents a single-industry solution but does not eliminate concentration.")
heading(doc, "3.2 Empirical-Bayes mean shrinkage", 2)
body(doc, "Each industry's sample mean is shrunk toward the cross-sectional grand mean. Let m_i be the sample mean, v_i/T its sampling variance, m_0 the grand mean and tau^2 the estimated between-industry prior variance. The posterior approximation is:")
equation(doc, "mu_i^B = B_i m_i + (1-B_i)m_0,    B_i = tau^2 / (tau^2 + v_i/T)")
body(doc, "Noisy industries receive more shrinkage; more precise estimates retain more of their sample signal.")
heading(doc, "3.3 Ledoit-Wolf covariance shrinkage", 2)
equation(doc, "Sigma_LW = delta F + (1-delta) Sigma_hat")
body(doc, "The target F is a scaled identity in the implementation provided by scikit-learn, and the shrinkage intensity delta follows the Ledoit-Wolf formula. The resulting covariance is positive definite and better conditioned. Shrinkage MVO combines mu^B and Sigma_LW in the same constrained quadratic programme as sample MVO.")
callout(doc, "Comparison", "Sample and shrinkage MVO use the same objective and constraints, isolating the effect of regularised inputs.")
page_break(doc)

heading(doc, "4. Tail risk and distributional robustness")
heading(doc, "4.1 Historical mean-CVaR", 2)
body(doc, "For portfolio loss L_t = -w'r_t, CVaR is represented using an auxiliary threshold eta and non-negative excess-loss variables u_t. The portfolio minimises empirical 95% CVaR subject to an estimated-return floor equal to the empirical-Bayes expected return of the equal-weight portfolio.")
equation(doc, "min_{w,eta,u}  eta + [1/((1-alpha)T)] sum_t u_t")
equation(doc, "u_t >= -w'r_t - eta;  u_t >= 0;  (mu^B)'w >= (mu^B)'(1/N)")
body(doc, "The return floor prevents the optimiser from reducing tail risk by accepting a lower estimated return than the benchmark. SciPy/HiGHS solves the linear programme.")
heading(doc, "4.2 Wasserstein robust CVaR", 2)
body(doc, "Historical CVaR treats the empirical distribution as exact. The robust model instead considers all distributions inside a 1-Wasserstein ball of radius epsilon around it. With an L1 ground metric and linear portfolio loss, the worst-case CVaR adds an infinity-norm concentration penalty:")
equation(doc, "Robust CVaR = Empirical CVaR + [epsilon/(1-alpha)] ||w||_infinity")
body(doc, "For long-only weights, ||w||_infinity is linearised by introducing m with m >= w_i. The radius epsilon is chosen on the internal validation block from 0.0005, 0.001, 0.002 and 0.004. Its median selected value was 0.002 monthly-return units.")
heading(doc, "4.3 Implementation and validation", 2)
for text in [
    "Python package with separate data, estimation, optimisation, backtest, metrics and plotting modules.",
    "Quadratic programmes solved by SLSQP; CVaR programmes solved by HiGHS linear programming.",
    "Unit tests check data continuity, covariance positive semidefiniteness, full investment, long-only bounds and the robust concentration penalty.",
    "All 129 monthly fits completed without an optimiser failure.",
]: bullet(doc, text)
page_break(doc)

heading(doc, "5. Main out-of-sample results")
result_rows = []
for strategy in summary.index:
    r = summary.loc[strategy]
    result_rows.append((
        strategy,
        pct(r["Annualized Return"], 1),
        pct(r["Annualized Volatility"], 1),
        num(r["Zero-Rate Sharpe"], 2),
        pct(r["Maximum Drawdown"], 1),
        pct(r["95% CVaR Loss"], 1),
    ))
table(doc, ["Strategy", "Return", "Vol.", "Sharpe", "Max DD", "CVaR loss"], result_rows,
      widths=[2.30, 0.82, 0.82, 0.78, 0.86, 0.95], font_size=7.6)
figure(doc, FIG / "cumulative_wealth.png", 6.55, "Figure 1. Growth of one dollar after monthly turnover costs, January 2015-September 2025.")
body(doc, "Equal weight finished with terminal wealth of 3.20 and the highest compound return. Sample MVO finished at 2.21 with slightly more annualised volatility than 1/N. Shrinkage MVO had the lowest volatility and best maximum drawdown. Both CVaR strategies improved risk control relative to sample MVO while retaining more upside than shrinkage MVO.", after=2.5)
callout(doc, "Result", "Controlling estimation error and the risk objective mattered more than maximising the in-sample mean-variance criterion.")
page_break(doc)

heading(doc, "6. Downside risk and economic behaviour")
figure(doc, FIG / "drawdowns.png", 6.55, "Figure 2. Drawdown paths based on net monthly returns.")
body(doc, f"Sample MVO experienced the deepest drawdown at {pct(summary.loc['Sample MVO','Maximum Drawdown'],1)}. Shrinkage MVO limited its maximum drawdown to {pct(summary.loc['Shrinkage MVO','Maximum Drawdown'],1)}, while historical CVaR reached {pct(summary.loc['Historical Mean-CVaR','Maximum Drawdown'],1)}. The estimated 95% monthly CVaR loss fell from {pct(summary.loc['Equal Weight','95% CVaR Loss'],1)} for equal weight to {pct(summary.loc['Shrinkage MVO','95% CVaR Loss'],1)} for shrinkage MVO and {pct(summary.loc['Historical Mean-CVaR','95% CVaR Loss'],1)} for historical CVaR.")
heading(doc, "6.1 Stress episodes", 2)
def stress_value(period, strategy, col="Cumulative Return"):
    return stress[(stress["Period"] == period) & (stress["Strategy"] == strategy)][col].iloc[0]
period_covid = "COVID sell-off (2020-02 to 2020-03)"
period_2022 = "2022 tightening sell-off (2022-01 to 2022-10)"
stress_rows = []
for strategy in summary.index:
    stress_rows.append((strategy, pct(stress_value(period_covid, strategy), 1), pct(stress_value(period_2022, strategy), 1)))
table(doc, ["Strategy", "COVID sell-off", "2022 tightening"], stress_rows, widths=[3.1, 1.75, 1.8], font_size=8.1)
body(doc, "Historical CVaR lost 18.0% in the February-March 2020 sell-off, compared with 22.7% for 1/N. During the 2022 tightening episode, shrinkage MVO performed best (-8.2%) and robust CVaR lost 18.8%. Distributional robustness does not protect against every regime.")
page_break(doc)

heading(doc, "7. Turnover, concentration and regime dependence")
figure(doc, FIG / "turnover.png", 6.05, "Figure 3. Average monthly one-way turnover. Ten basis points are charged per unit of turnover.")
trade_rows = []
for strategy in summary.index:
    r = summary.loc[strategy]
    trade_rows.append((strategy, pct(r["Average Monthly Turnover"],1), pct(r["Annualized Cost Drag"],2), num(r["Average Effective N"],1), pct(r["Average Maximum Weight"],1)))
table(doc, ["Strategy", "Turnover", "Cost drag", "Effective N", "Avg. max wt."], trade_rows,
      widths=[2.45, 0.95, 0.95, 1.0, 1.15], font_size=7.7)
body(doc, "Sample MVO reallocated 22.0% per month on average and held an effective 3.5 industries. Its average largest position was close to the 35% cap. Wasserstein robustness increased effective breadth to 4.8 industries and reduced the average largest position to 24.4% because the L1 ambiguity penalty directly penalises the largest weight.")
heading(doc, "7.1 Volatility regimes", 2)
regime_rows = []
for strategy in summary.index:
    high = regimes[(regimes.Strategy == strategy) & (regimes.Regime == "High volatility")].iloc[0]
    low = regimes[(regimes.Strategy == strategy) & (regimes.Regime == "Lower volatility")].iloc[0]
    regime_rows.append((strategy, num(high["Zero-Rate Sharpe"],2), num(low["Zero-Rate Sharpe"],2)))
table(doc, ["Strategy", "High-vol. Sharpe", "Lower-vol. Sharpe"], regime_rows, widths=[3.1, 1.75, 1.8], font_size=8.1)
body(doc, "Historical CVaR had the strongest high-volatility Sharpe (0.91), narrowly above equal weight (0.90). Shrinkage MVO performed better in lower-volatility months (0.78) than in high-volatility months (0.67). Sample MVO's lower-volatility Sharpe was only 0.24, consistent with unstable mean estimates rather than compensation for systematic stress exposure.")
page_break(doc)

heading(doc, "8. Statistical inference, robustness and limitations")
body(doc, "A 2,000-draw circular moving-block bootstrap with six-month blocks preserves short-run dependence. Each optimisation strategy is compared with equal weight using annualised arithmetic mean-return and zero-rate Sharpe differences. Confidence intervals are percentile intervals.")
boot_rows = []
for strategy, r in bootstrap.iterrows():
    boot_rows.append((
        strategy,
        pct(r["Annualized Mean Difference"],1),
        f"[{pct(r['Mean Difference CI Low'],1)}, {pct(r['Mean Difference CI High'],1)}]",
        num(r["Sharpe Difference"],2),
        f"[{num(r['Sharpe Difference CI Low'],2)}, {num(r['Sharpe Difference CI High'],2)}]",
    ))
table(doc, ["Strategy vs 1/N", "Mean diff.", "95% CI", "Sharpe diff.", "95% CI"], boot_rows,
      widths=[2.2, 0.9, 1.45, 0.95, 1.35], font_size=7.5)
body(doc, "Every mean-return difference was negative. The confidence intervals for historical and robust CVaR included zero, while the shrinkage MVO interval ended close to zero. No Sharpe-difference interval showed a positive advantage. Robust CVaR's observed Sharpe was 0.008 below 1/N, with a 47.3% bootstrap probability of a positive difference. This supports non-superiority, not equality.")
heading(doc, "8.1 Limitations", 2)
for text in [
    "Industry portfolios are research constructs rather than directly investable securities; implementation costs for actual products may differ.",
    "The universe excludes bonds, cash and international assets, so results are not a complete strategic asset-allocation study.",
    "Sixty monthly observations leave only a few empirical observations in the 5% CVaR tail, increasing sampling uncertainty.",
    "The Wasserstein radius is validation-selected rather than calibrated as a formal confidence radius.",
    "A zero risk-free rate is used for the labelled Sharpe statistic, and inference is sensitive to block length and sample length.",
    "Historical backtests cannot establish future profitability; model and data vintages may change.",
]: bullet(doc, text, size=8.9)
callout(doc, "Reproducibility", "The repository stores monthly weights, returns, turnover, selected hyperparameters, bootstrap results and figures.")
page_break(doc)

heading(doc, "9. Conclusion")
body(doc, "Sample MVO performed poorly because noisy means and covariances produced concentrated, high-turnover decisions. Empirical-Bayes and Ledoit-Wolf shrinkage improved realised risk. CVaR shifted the objective from variance to tail losses and performed better in the COVID sell-off.")
body(doc, "Wasserstein robustness produced a broader portfolio than historical CVaR and the strongest Sharpe among the optimisation methods, but it did not beat equal weighting on compound return or statistical inference. The robust penalty therefore helped control concentration and distributional uncertainty without solving the more fundamental problem of forecasting expected returns. Equal weighting remained the strongest overall benchmark.")
body(doc, "Financial models should separate estimation, decision and evaluation. A sophisticated objective has value only if its decisions hold up on unseen data, after costs and under stress.")
heading(doc, "References", 2)
refs = [
    "DeMiguel, V., Garlappi, L. and Uppal, R. (2009). Optimal Versus Naive Diversification: How Inefficient Is the 1/N Portfolio Strategy? Review of Financial Studies, 22(5), 1915-1953. doi:10.1093/rfs/hhm075.",
    "French, K. R. (2025 data vintage). 12 Industry Portfolios. Kenneth R. French Data Library, Dartmouth College.",
    "Ledoit, O. and Wolf, M. (2004). A Well-Conditioned Estimator for Large-Dimensional Covariance Matrices. Journal of Multivariate Analysis, 88(2), 365-411. doi:10.1016/S0047-259X(03)00096-4.",
    "Markowitz, H. (1952). Portfolio Selection. Journal of Finance, 7(1), 77-91. doi:10.1111/j.1540-6261.1952.tb01525.x.",
    "Mohajerin Esfahani, P. and Kuhn, D. (2018). Data-Driven Distributionally Robust Optimization Using the Wasserstein Metric. Mathematical Programming, 171, 115-166. doi:10.1007/s10107-017-1172-1.",
    "Rockafellar, R. T. and Uryasev, S. (2000). Optimization of Conditional Value-at-Risk. Journal of Risk, 2(3), 21-41. doi:10.21314/JOR.2000.038.",
]
for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.20)
    p.paragraph_format.first_line_indent = Inches(-0.20)
    p.paragraph_format.space_after = Pt(2.2)
    p.paragraph_format.line_spacing = 0.98
    set_font(p.add_run(ref), size=8.5)
heading(doc, "Repository deliverables", 2)
body(doc, "The accompanying repository includes the versioned data snapshot, official-data downloader, tested Python package, walk-forward runner, all output CSV files, figures, editable Word report and PDF report.", after=0, size=8.9)

doc.core_properties.title = "Estimation Error in Portfolio Optimisation"
doc.core_properties.subject = "Bayesian shrinkage, CVaR and Wasserstein robust portfolio construction"
doc.core_properties.author = "Ben Heskin"
doc.core_properties.keywords = "portfolio optimisation, Bayesian shrinkage, CVaR, Wasserstein, distributionally robust optimisation"
OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT_DOCX)
print(OUT_DOCX)
